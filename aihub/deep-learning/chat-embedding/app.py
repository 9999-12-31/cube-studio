
import subprocess
subprocess.run('pip install openai html2text', shell=True, capture_output=True, text=True)

import io,sys,os,base64,pysnooper

import pandas as pd
from cubestudio.aihub.model import Model,Validator,Field_type,Field
import numpy,pandas
import os,shutil,csv
from scipy.spatial.distance import cdist
import datetime,copy,requests,re,json
from flask import g
from modelscope.pipelines import pipeline
from modelscope.utils.constant import Tasks


class Chat_embedding_Model(Model):
    # 模型基础信息定义
    name='chat-embedding'   # 该名称与目录名必须一样，小写
    label='gpt私有知识库'
    describe="gpt私有知识库"
    field="自然语言"  # [机器视觉，听觉，自然语言，多模态，大模型]
    scenes="聊天机器人"
    status='online'
    version='v20221001'
    pic='example.jpeg'  # 离线图片，作为模型的样式图，330*180尺寸比例
    images = 'ccr.ccs.tencentyun.com/cube-studio/modelscope:base-cuda11.3-python3.7'
    # 和train函数的输入参数对应，并且会对接显示到pipeline的模板参数中
    # 和inference函数的输入参数对应，并且会对接显示到web界面上
    inference_inputs = [
        Field(Field_type.text, name='chat_id', label='场景名', describe='可以配置不同场景的名称来做细分召回',
              default='cube-studio', validators=Validator(max=200)),
        Field(Field_type.text, name='query', label='你的问题', describe='你的问题，最长200字',default='如何部署cube-studio平台？',validators=Validator(max=200)),
    ]
    inference_outputs = ['markdown' for x in range(20)]

    # 会显示在web界面上，让用户作为示例输入
    web_examples=[
        {
            "label": "示例一描述",
            "input": {
                "query": '如何部署cube-studio平台？',
                "chat_id": "cube-studio"
            }
        }
    ]
    all_emb = {}
    all_keyword = {}
    all_summary = {}
    # 转向量
    # @pysnooper.snoop()
    def to_vector(self,text_arr):
        result = self.embedding_model(input={"source_sentence": text_arr})
        embedding = result['text_embedding']
        print(embedding.shape)
        return embedding

    # 获取摘要
    # @pysnooper.snoop()
    def to_summary(self,text):
        if text.strip():
            # 最大长度 4096个token
            contents = self.simple_content_split_max(text=text,max_token=2048)
            file_summary = ''
            for content in contents:
                summary=''
                # try:
                #     summary = get_entry(content.strip())
                # except Exception as e:
                #     print(e)
                if not summary:
                    try:
                        # 模型效果太差
                        result = self.summary_model(content)
                        summary = result['text']
                    except Exception as e:
                        print(e)

                file_summary+="，"+summary
            return file_summary
        else:
            return ''

    # 获取摘要
    # @pysnooper.snoop()
    def to_keyword(self,text):
        if text.strip():
            contents = self.simple_content_split_max(text, max_token=500)
            keywords = []
            for content in contents:
                if content.strip():
                    # 先用llm
                    keyword=[]
                    # try:
                    #     keyword = get_entry(content.strip())
                    # except Exception as e:
                    #     print(e)
                    if not keyword:
                        # ner 模型提取的太少了，不好。直接用llm
                        try:
                            result = self.ner_model(content)
                            keyword = [x['span'] for x in result['output']]
                        except Exception as e:
                            print(e)
                    keywords.extend(keyword)
            keywords = list(set(keywords))

            return keywords
        else:
            return []


    # 话题分割
    # @pysnooper.snoop()
    def topic_segmentation(self,text):
        result = self.topic_seg_model(text)
        result = result['text'].split('\n')
        return result

    # 对长文本简单分割,尽量每一块能有max_token大小
    def simple_content_split_max(self,text, max_token=500):
        separators = ['\n\n','```', '\n', '。', '，', ' ']
        paragraphs = [text]

        for sep in separators:

            if all(len(p) <= max_token for p in paragraphs):
                break
            # 如果存在长度长多max_token，就尝试使用分割
            else:
                temp_paragraphs = []
                for paragraph in paragraphs:
                    if len(paragraph) <= max_token:
                        temp_paragraphs.append(paragraph)
                    else:
                        sub_paragraphs = paragraph.split(sep)
                        # 尽量合并过短的片段
                        insert_str=''
                        for sub in sub_paragraphs:
                            if len(insert_str+sub)>max_token:
                                if insert_str.strip():
                                    temp_paragraphs.append(sub)
                                insert_str=sub
                            else:
                                insert_str+=sub
                        # 插入最后一个拼接字符串
                        if insert_str.strip():
                            temp_paragraphs.append(insert_str)
                paragraphs = temp_paragraphs


        # 再检验一遍还有没切完的，直接长度硬分割
        all_paragraphs = []
        for p in paragraphs:
            if len(p)>max_token:
                arr = [p[i:i + max_token] for i in range(0, len(p), max_token)]
                all_paragraphs.extend(arr)
            else:
                all_paragraphs.append(p)

        return all_paragraphs

    # 对长文本简单分割,遇到分割符就完全分割
    def simple_content_split_min(self,text, max_token=500):
        separators = ['\n\n', '\n', '。', '，', ' ']
        paragraphs = [text]

        for sep in separators:

            if all(len(p) <= max_token for p in paragraphs):
                break
            # 如果存在长度长多max_token，就尝试使用分割
            else:
                temp_paragraphs = []
                for paragraph in paragraphs:
                    if len(paragraph) <= max_token:
                        temp_paragraphs.append(paragraph)
                    else:
                        sub_paragraphs = paragraph.split(sep)
                        temp_paragraphs.extend(sub_paragraphs)
                paragraphs = temp_paragraphs

        # 再检验一遍还有没切完的，直接长度硬分割
        all_paragraphs = []
        for p in paragraphs:
            if len(p)>max_token:
                arr = [p[i:i + max_token] for i in range(0, len(p), max_token)]
                all_paragraphs.extend(arr)
            else:
                all_paragraphs.append(p)

        return all_paragraphs

    # 自定义文件分割
    def custom_content_split(self,content,**kwargs):
        split_type=kwargs.get('split_type','\n')
        if split_type=='\n':
            return content.split("\n")
        if split_type=='re':
            re_str = kwargs.get('re')
            return re.split(re_str,content)
        if split_type=='re+summary':
            re_str = kwargs.get('re')
            src_content = re.split(re_str, content)

    # es召回
    def es_recall(self):
        pass

    # 插入es
    def es_insert(self,chat_id,data):
        data=[
            {
                "col":"question",
                "value":"aaaaa",
                "weight":"0.1"
            }
        ]
        # data需要是[]
        pass

    # 向量召回
    # @pysnooper.snoop()
    def vector_recall(self,embedding,embedding_arr,topn=20):
        print(embedding.shape)
        print(embedding_arr.shape)
        dist = cdist(embedding, embedding_arr, metric='euclidean')
        distances = dist.tolist()[0]
        dist = pd.DataFrame(dist).T
        dist.columns = ['score']
        index = dist.sort_values(['score'],ascending=True).iloc[0:topn,:].index.tolist()
        scores = [distances[i] for i in index]
        return index,scores

    # 排序算法
    # @pysnooper.snoop()
    def ranking(self,source_sentence,sentences_to_compare):
        if type(sentences_to_compare) == str:
            sentences_to_compare = sentences_to_compare.split('\n')
        sentences_to_compare = [x.strip() for x in sentences_to_compare if x.strip()]
        result = self.rank_model(input={"source_sentence": [source_sentence], "sentences_to_compare": sentences_to_compare})

        # 按相似度排序
        sim = []
        for index,compare in enumerate(sentences_to_compare):
            sim.append([compare,round(result['scores'][index],2)])
        sim = sorted(sim,key=lambda x:x[1],reverse=True)
        return sim

    # 读取文件整体内容
    def read_file(self,file_path):
        # pdf文件
        if '.pdf' in file_path:
            from load_pdf import load_pdf
            text = load_pdf(file_path)
            return text

        # doc文件
        if '.doc' in file_path:
            import docx

            doc = docx.Document(file_path)
            text = ''
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n\n"
                print(paragraph.text)
            return text


        # 文本文件，自动按行分端，
        if '.txt' in file_path:
            text = open(file_path, mode='r').read()
            return text

        # 网页文件
        if '.html' in file_path or '.htm' in file_path:
            from load_html import load_html
            content = load_html(file_path)
            return content

        # markdown
        if '.md' in file_path.lower():
            return open(file_path, mode='r').read()

        return open(file_path, mode='r').read()

    # 加载单个文件
    # @pysnooper.snoop()
    def seg_file(self,file_path):
        """
        读取单个文件
        @param chat: 场景对象
        @param file_path: 单个文件地址
        @return: 返回文本数组
        """
        # pdf文件
        if '.pdf' in file_path:
            from load_pdf import load_pdf
            text = load_pdf(file_path)
            text = self.topic_segmentation(text)
            return text

        # doc文件
        if '.doc' in file_path:
            import docx

            doc = docx.Document(file_path)
            text = ''
            for paragraph in doc.paragraphs:
                text +=paragraph.text+"\n\n"
                print(paragraph.text)
            text = self.topic_segmentation(text)
            return text

        # csv文件,csv文件自动按行分端
        if '.csv' in file_path:
            all_content = ['\n'.join(f"{k.strip()}: {v.strip()}" for k, v in row.items() if k and v) for row in csv.DictReader(open(file_path,encoding="utf-8-sig"))]
            return all_content

        # 文本文件，自动按行分端，
        if '.txt' in file_path:
            text = open(file_path, mode='r').readlines()
            text = [x.strip() for x in text if x.strip()]
            return text
            # text = ''.join(open(file_path, mode='r').readlines())
            # text = self.topic_segmentation(text)
            # return text

        # 网页文件
        if '.html' in file_path or '.htm' in file_path:
            from load_html import load_html
            from load_markdown import load_markdown, markdown2segment
            all_content = load_html(file_path)
            file_name = os.path.splitext(os.path.basename(file_path))[0]
            contents = markdown2segment(titles=[file_name], markdown=all_content)
            return contents

        # markdown
        if '.md' in file_path.lower():
            from load_markdown import load_markdown,markdown2segment
            all_content = load_markdown(file_path)
            file_name = os.path.splitext(os.path.basename(file_path))[0]
            contents = markdown2segment(titles=[file_name], markdown=all_content)

            return contents

        text = open(file_path, mode='r').readlines()
        text = [x.strip() for x in text if x.strip()]
        return text

    # 设置多级目录，多次召回
    def toc(self):
        pass

    def recall(self,knowledge_base_id,question,history=[],**kwargs):
        back = self.topk(chat_id=knowledge_base_id,query=question,**kwargs)
        return back

    @pysnooper.snoop()
    def topk(self,chat_id, query, topn=20, min_Score=0, max_token=3800):
        if chat_id not in self.all_emb:
            return []
        # 向量召回
        embedding = self.to_vector([query])

        # embedding文本块召回
        index_arr,score_arr = self.vector_recall(embedding,self.all_emb[chat_id].get('embeddings',[]),topn=max(10,topn))
        text_arr = [self.all_emb[chat_id]['texts'][i] for i in index_arr]
        file_arr = [self.all_emb[chat_id]['files'][i] for i in index_arr]

        # 排序
        result = self.ranking(query,text_arr)
        back = []
        for text,score in result[:topn]:
            file=''
            if text in text_arr:
                file = file_arr[text_arr.index(text)]
            if score>min_Score:
                back.append({
                    "context":text,
                    "score":score,
                    "file":file
                })
        if not back:
            back = [{
                "context":result[0][0],
                "score": result[0][1],
                "file":file_arr[0],
            }]
        print(back)
        return back

    # 加载模型，所有一次性的初始化工作可以放到该方法下。注意save_model_dir必须和训练函数导出的模型结构对应
    def load_model(self,save_model_dir=None,**kwargs):
        self.embedding_model = pipeline('sentence-embedding', 'damo/nlp_corom_sentence-embedding_chinese-base')
        self.rank_model = pipeline('text-ranking', 'damo/nlp_corom_passage-ranking_chinese-tiny')
        # self.topic_seg_model = pipeline('document-segmentation', 'damo/nlp_bert_document-segmentation_chinese-base')
        # self.ner_model = pipeline('named-entity-recognition','damo/nlp_structbert_keyphrase-extraction_base-icassp2023-mug-track4-baseline')
        # self.summary_model = pipeline('extractive-summarization', 'damo/nlp_ponet_extractive-summarization_doc-level_chinese-base')

        # 加载索引
        self.all_emb ={}
        if os.path.exists('embedding/all_embedding.json'):
            save_embedding = json.load(open('embedding/all_embedding.json', mode='r'))
            for chat_id in save_embedding:
                save_path = f'embedding/{chat_id}.npy'
                if os.path.exists(save_path):
                    embeddings = numpy.load(save_path)
                    self.all_emb[chat_id]={
                        "files":save_embedding[chat_id]['files'],
                        "texts": save_embedding[chat_id]["texts"],
                        "embeddings":embeddings
                    }

        # 加载关键词
        self.all_keyword = {}
        if os.path.exists('keyword/all_keyword.json'):
            self.all_keyword = json.load(open('keyword/all_keyword.json',mode='r'))

        # 加载总结
        if os.path.exists('summary/all_summary.json'):
            save_embedding = json.load(open('summary/all_summary.json', mode='r'))
            for chat_id in save_embedding:
                save_path = f'summary/{chat_id}.npy'
                if os.path.exists(save_path):
                    embeddings = numpy.load(save_path)
                    self.all_summary[chat_id] = {
                        "files": save_embedding[chat_id]['files'],
                        "summarys": save_embedding[chat_id]["summarys"],
                        "embeddings": embeddings
                    }

    def get_file_paths(self,directory):
        file_paths = []
        for root, dirs, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                if '.md' in file_path or '.html' in file_path or '.csv' in file_path or '.txt' in file_path or '.pdf' in file_path or '.doc' in file_path:
                    file_paths.append(file_path)
        return file_paths

    # embedding 文本
    # 初始化文本成embedding
    # @pysnooper.snoop()
    @pysnooper.snoop()
    def seg_emb_init(self,chat_id,files):

        all_file_content = []

        index = 0
        for file in files:
            print(f'加载{file}')
            contents = self.seg_file(file)
            for content in contents:
                new_contents=[]
                if len(content) > 1000:
                    print(f'文件{file},中存在分片内容过长,进行再分割')
                    # 可以使用主题分割
                    new_contents = self.simple_content_split_max(content,max_token=500)
                if new_contents:
                    for sub_content in new_contents:
                        if sub_content.strip():
                            all_file_content.append({
                                "id": str(index),
                                "file": file,
                                "text": sub_content,
                                "len": len(sub_content)
                            })
                            index+=1
                elif content.strip():
                    all_file_content.append({
                        "id":str(index),
                        "file":file,
                        "text":content,
                        "len":len(content)
                    })
                    index+=1

        # print(all_content)
        all_content = [x['text'] for x in all_file_content]
        all_files = [x['file'] for x in all_file_content]
        print('检索到分片文本%s个'%len(all_content))

        embeddings = self.to_vector(all_content)

        # embedding必须要矩阵存储，这样计算才快
        print(embeddings.shape)
        self.all_emb[chat_id]={
            "files":all_files,
            "texts":all_content,
            "embeddings":embeddings
        }
        # 数据保存
        os.makedirs('embedding', exist_ok=True)
        # 保存索引
        save_path = f'embedding/{chat_id}.npy'
        if os.path.exists(save_path):
            os.remove(save_path)
        numpy.save(save_path, embeddings)

        # 保存到文件
        save_embedding=json.load(open('embedding/all_embedding.json')) if os.path.exists('embedding/all_embedding.json') else {}
        save_embedding[chat_id]={
            "files":all_files,
            "texts":all_content,
            "embeddings":save_path
        }

        json.dump(save_embedding,open('embedding/all_embedding.json',mode='w'),indent=4,ensure_ascii=False)

        return {
            "files":files,
            "texts":all_content,
            "embeddings":embeddings
        }


    # 文章整体 关键词提取
    # @pysnooper.snoop()
    def keyword_init(self, chat_id, files):
        all_files_keywords=[]
        for file in files:
            content = self.read_file(file)
            if content.strip():
                keywords = self.to_keyword(content)
                all_files_keywords.append(
                    {
                        "file": file,
                        "keywords": keywords,
                    }
                )

        # 去除了空白文件的
        files = [x['file'] for x in all_files_keywords]
        all_content = [x['keywords'] for x in all_files_keywords]
        self.all_keyword[chat_id]={
            "files":files,
            "keywords":all_content
        }

        # 数据保存
        os.makedirs('keyword', exist_ok=True)
        json.dump(self.all_keyword,open('keyword/all_keyword.json',mode='w'),indent=4,ensure_ascii=False)

        # 关键词放es
        return all_files_keywords


    # 文章整体概要提取
    # @pysnooper.snoop()
    def summary_emb_init(self, chat_id, files):
        all_files_summary=[]
        for file in files:
            content = self.read_file(file)
            if content.strip():
                summary = self.to_summary(content)
                all_files_summary.append(
                    {
                        "file": file,
                        "summary": summary
                    }
                )

        # 总结要做embedding
        # print(all_content)
        files = [x['file'] for x in all_files_summary]
        all_content = [x['summary'] for x in all_files_summary]
        print('总结文本%s个' % len(all_content))

        embeddings = self.to_vector(all_content)

        # embedding必须要矩阵存储，这样计算才快
        print(embeddings.shape)
        self.all_summary[chat_id] = {
            "files":files,
            "summarys":all_content,
            "embeddings":embeddings
        }

        # 保存embedding
        os.makedirs('summary', exist_ok=True)

        save_path = f'summary/{chat_id}.npy'
        if os.path.exists(save_path):
            os.remove(save_path)
        numpy.save(save_path, embeddings)

        # 保存到文件
        save_embedding=json.load(open('summary/all_summary.json')) if os.path.exists('summary/all_summary.json') else {}
        save_embedding[chat_id]={
            "files":files,
            "summarys":all_content,
            "embeddings":save_path
        }

        json.dump(save_embedding,open('summary/all_summary.json',mode='w'),indent=4,ensure_ascii=False)

        return {
            "files":files,
            "summarys":all_content,
            "embeddings":embeddings
        }

    # @pysnooper.snoop()
    def file_init(self,chat_id,knowledge_config):
        # 读取文件列表
        files = knowledge_config.get('file',[])

        if type(files) != list:
            files = [files]
        new_file_list = []
        for file in files:
            # 如何使用目录，就直接遍历
            if os.path.isdir(file):
                file_list = self.get_file_paths(file)
                new_file_list += file_list
        files = [file for file in files if os.path.isfile(file) and os.path.exists(file)] + new_file_list
        print('检索到文件 %s个' % len(files))

        # 分割和embedding
        self.seg_emb_init(chat_id=chat_id, files=files)

        # # 关键词，和强化词
        # self.keyword_init(chat_id=chat_id, files=files)
        #
        # # # summary和embedding
        # self.summary_emb_init(chat_id=chat_id, files=files)

    # 上传文件
    # @pysnooper.snoop()
    def upload_files(self,chat_id,files,**kwargs):
        print(kwargs)
        print(files)
        files_path=[]
        if type(files)==dict:
            files = list(files.values())
        for file in files:
            file_name = file.filename
            print(file_name)
            save_path = f'/data/k8s/kubeflow/global/knowledge/{file_name}'
            os.makedirs(os.path.dirname(save_path),exist_ok=True)
            file.save(save_path)
            files_path.append(save_path)
        knowledge_config={
            "type": "file",
            "file":files_path
        }
        self.file_init(chat_id=chat_id, knowledge_config=knowledge_config)

    # 单独做文件分割的，用于提供成接口
    def files_seg(self,files,**kwargs):
        files_path=[]
        if type(files)==dict:
            files = list(files.values())
        for file in files:
            file_name = file.filename
            print(file_name)
            save_path = f'/data/k8s/kubeflow/global/knowledge/linshi/{file_name}'
            os.makedirs(os.path.dirname(save_path),exist_ok=True)
            file.save(save_path)
            files_path.append(save_path)

        all_file_content=[]
        index=0
        for file in files_path:
            print(f'加载{file}')
            contents = self.seg_file(file)
            for content in contents:
                new_contents=[]
                if len(content) > 500:
                    print(f'文件{file},中存在分片内容过长,进行再分割')
                    # 可以使用主题分割
                    new_contents = self.simple_content_split_max(content,max_token=500)
                if new_contents:
                    for sub_content in new_contents:
                        if sub_content.strip():
                            all_file_content.append({
                                "id": str(index),
                                "file": file,
                                "text": sub_content,
                                "len": len(sub_content)
                            })
                            index+=1
                elif content.strip():
                    all_file_content.append({
                        "id":str(index),
                        "file":file,
                        "text":content,
                        "len":len(content)
                    })
                    index+=1
        return all_file_content

    # web每次用户请求推理，用于对接web界面请求
    @pysnooper.snoop()
    def inference(self,query,chat_id='cube-studio',knowledge_config=None,**kwargs):
        if knowledge_config:
            self.file_init(chat_id=chat_id, knowledge_config=knowledge_config)
        knowledge = self.topk(chat_id=chat_id, query=query)
        # knowledge = [doc['context'] for doc in knowledge]
        # knowledge = '\n\n-------------\n\n'.join(knowledge)

        back=[
            {
                "markdown":doc.get('context','')+"\n\n得分："+str(doc.get('score',''))+"\n\n文件："+doc.get('file','')+"\n\n\n",
            } for doc in knowledge
        ]
        return back

model=Chat_embedding_Model()

# 容器中运行调试推理时
model.load_model(save_model_dir=None)

result = model.inference(query='如何部署cube-studio平台？',chat_id='cube-studio', knowledge_config={
    "type": "file",
    "file":"cube-studio.csv"
})
print(result)

# 1、基础信息如何加入到文本片段中，比如这个文档是关于cube-studio的，也就是对于文件来说的普适性知识
# 2、

# 模型启动web时使用
if __name__=='__main__':
    model.run()

